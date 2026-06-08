"""
Document loading and text extraction.

Supports: PDF (pypdf), DOCX (python-docx), TXT (plain).
Each loader returns a list of LangChain Document objects with rich metadata.
"""
from __future__ import annotations

import hashlib
import io
from pathlib import Path
from typing import TYPE_CHECKING

import structlog
from langchain_core.documents import Document

if TYPE_CHECKING:
    from fastapi import UploadFile

log = structlog.get_logger()


# ── helpers ───────────────────────────────────────────────────────────────────

def _file_hash(data: bytes) -> str:
    return hashlib.sha256(data).hexdigest()[:16]


def _base_meta(filename: str, file_hash: str) -> dict:
    return {"source": filename, "file_hash": file_hash}


# ── loaders ───────────────────────────────────────────────────────────────────

def load_pdf(data: bytes, filename: str) -> list[Document]:
    from pypdf import PdfReader

    fhash = _file_hash(data)
    reader = PdfReader(io.BytesIO(data))
    docs: list[Document] = []

    for i, page in enumerate(reader.pages):
        text = page.extract_text() or ""
        text = text.strip()
        if not text:
            continue
        meta = {**_base_meta(filename, fhash), "page": i + 1, "total_pages": len(reader.pages)}
        docs.append(Document(page_content=text, metadata=meta))

    log.info("pdf_loaded", filename=filename, pages=len(docs))
    return docs


def load_docx(data: bytes, filename: str) -> list[Document]:
    import docx

    fhash = _file_hash(data)
    doc = docx.Document(io.BytesIO(data))
    full_text = "\n\n".join(p.text for p in doc.paragraphs if p.text.strip())
    meta = {**_base_meta(filename, fhash), "paragraphs": len(doc.paragraphs)}

    log.info("docx_loaded", filename=filename)
    return [Document(page_content=full_text, metadata=meta)]


def load_txt(data: bytes, filename: str) -> list[Document]:
    fhash = _file_hash(data)
    text = data.decode("utf-8", errors="replace").strip()
    meta = {**_base_meta(filename, fhash), "chars": len(text)}
    log.info("txt_loaded", filename=filename)
    return [Document(page_content=text, metadata=meta)]


# ── dispatcher ────────────────────────────────────────────────────────────────

LOADERS = {
    ".pdf": load_pdf,
    ".docx": load_docx,
    ".txt": load_txt,
}


def load_document(data: bytes, filename: str) -> list[Document]:
    """Dispatch to the right loader based on file extension."""
    ext = Path(filename).suffix.lower()
    loader = LOADERS.get(ext)
    if loader is None:
        raise ValueError(f"Unsupported file type: {ext!r}. Supported: {list(LOADERS)}")
    return loader(data, filename)


async def load_upload(file: "UploadFile") -> list[Document]:
    """Load a FastAPI UploadFile and return Document list."""
    data = await file.read()
    return load_document(data, file.filename or "unknown")
